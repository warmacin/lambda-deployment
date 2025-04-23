import boto3, datetime, os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# CloudWatch client
cw = boto3.client('cloudwatch', region_name='ap-south-1')
s3 = boto3.client('s3')

# Your single instance
INSTANCE_ID = "i-0bf7a1cdc65f2a2fe"

# Metrics to include (name, label)
METRICS = [
    ("NetworkIn",         "Network In"),
    ("NetworkOut",        "Network Out"),
    ("mem_used_percent",  "Memory Utilization"),
    ("cpu_usage_idle",    "CPU Idle %"),
]

def check_alarms(metric_name):
    resp = cw.describe_alarms_for_metric(
        MetricName=metric_name,
        Namespace="AWS/EC2",
        Dimensions=[{"Name":"InstanceId","Value":INSTANCE_ID}]
    )
    alarms = resp.get("MetricAlarms", [])
    return [a["AlarmName"] for a in alarms] or ["No alarms"]

def fetch_chart(metric_name, title):
    widget = {
      "metrics": [["AWS/EC2", metric_name, "InstanceId", INSTANCE_ID]],
      "start": "-PT24H", "end": "P0D",
      "view":"timeSeries","stat":"Average",
      "width":800, "height":300, "title": title
    }
    img = cw.get_metric_widget_image(MetricWidget=json.dumps(widget))["MetricWidgetImage"]
    return img

def lambda_handler(event, context):
    doc = Document()
    # Title
    h = doc.add_heading("SCL Resource Utilization Report", level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h.runs[0].font.size = Pt(18)
    h.runs[0].font.bold  = True
    doc.add_paragraph(f"Date: {datetime.date.today():%d-%m-%Y}")\
       .alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    # Instance section
    doc.add_heading(f"Instance: {INSTANCE_ID}", level=2)
    # Alarms summary
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Alarms"
    for m,name in METRICS:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = ", ".join(check_alarms(m))
    doc.add_page_break()

    # 2×2 charts
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = 'Table Grid'
    coords = [(0,0),(0,1),(1,0),(1,1)]
    for (m,name), (r,c) in zip(METRICS, coords):
        cell = tbl.rows[r].cells[c]
        cell.paragraphs[0].add_run(name).bold = True
        img = fetch_chart(m, name)
        path = f"/tmp/{m}.png"
        with open(path,"wb") as f: f.write(img)
        cell.add_paragraph().add_run().add_picture(path, width=Inches(3.5))

    # Save & upload
    out = f"/tmp/SCL_Report_{datetime.date.today():%d-%m-%Y}.docx"
    doc.save(out)
    bucket = event.get("ReportBucket","test-resource-util")
    key    = f"Test/{datetime.date.today():%Y-%m-%d}/SCL_Report.docx"
    s3.upload_file(out, bucket, key)

    return {"statusCode":200, "body":f"Uploaded to s3://{bucket}/{key}"}
