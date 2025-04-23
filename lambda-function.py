import sys
import json
import boto3
import os
import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# AWS clients
cw = boto3.client("cloudwatch", region_name="ap-south-1")
s3 = boto3.client("s3", region_name="ap-south-1")

# Constants for S3 location of your logos
LOGO_BUCKET      = "test-resource-util"
LOGO_PREFIX      = "lambda/Logos"
LEFT_LOGO_NAME   = "Hathi-Cement.png"
RIGHT_LOGO_NAME  = "Sapphire.jpg"

INSTANCE_ID   = "i-0bf7a1cdc65f2a2fe"
INSTANCE_NAME = "PRD_S4H_SCL_APP"
REGION        = "ap-south-1"
METRICS = [
    ("NetworkIn",        "Network In"),
    ("NetworkOut",       "Network Out"),
    ("mem_used_percent", "Memory Utilization"),
    ("CPUUtilization",   "CPU Utilization")
]

def fetch_chart(metric, instance_label):
    widget = {
        "metrics": [["AWS/EC2", metric, "InstanceId", INSTANCE_ID]],
        "stat": "Average",
        "view": "timeSeries",
        "stacked": False,
        "region": REGION,
        "start": "-PT24H",
        "end": "P0D",
        "width": 1100,
        "height": 300,
        "title": instance_label,
        "setPeriodToTimeRange": True,
        "legend": {"position": "bottom"}
    }
    resp = cw.get_metric_widget_image(MetricWidget=json.dumps(widget))
    return resp["MetricWidgetImage"]

def download_logos_to_tmp():
    """
    Download both logos from S3 into /tmp so they can be used by python-docx.
    """
    for name in (LEFT_LOGO_NAME, RIGHT_LOGO_NAME):
        key = f"{LOGO_PREFIX}/{name}"
        dest = f"/tmp/{name}"
        s3.download_file(LOGO_BUCKET, key, dest)

def set_header_footer(doc):
    """
    Build a two-cell header table: left Hathi logo, right Sapphire logo.
    No centered text.
    """
    section = doc.sections[0]
    header = section.header

    # create a 1x2 table
    table = header.add_table(rows=1, cols=2)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # remove table borders
    tbl = table._tbl
    for cell in tbl.iter_tcs():
        cell.tcPr.append(OxmlElement("w:tcBorders"))

    # Left cell: Hathi logo
    left_cell = table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    left_para.add_run().add_picture(f"/tmp/{LEFT_LOGO_NAME}", width=Inches(1.0))

    # Right cell: Sapphire logo
    right_cell = table.rows[0].cells[1]
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    right_para.add_run().add_picture(f"/tmp/{RIGHT_LOGO_NAME}", width=Inches(1.5))

    # Footer unchanged
    footer = section.footer.paragraphs[0]
    footer.text = f"Generated on {datetime.date.today():%d-%m-%Y}"
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer.runs[0].font.size = Pt(8)

def lambda_handler(event, context):
    try:
        # first, pull down the logos so /tmp has them
        download_logos_to_tmp()

        today = datetime.date.today().strftime("%d-%m-%Y")
        bucket = event.get("ReportBucket", "test-resource-util")
        key = f"Test/{datetime.date.today():%Y-%m-%d}/SCL_Report.docx"

        doc = Document()

        # use the new header/footer setup
        set_header_footer(doc)

        # rest of your report as before…
        title = doc.add_paragraph()
        run = title.add_run("SCL Resource Utilization")
        run.font.size = Pt(20)
        run.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        date_para = doc.add_paragraph(f"Date: {today}")
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()

        inst = doc.add_paragraph()
        inst_run = inst.add_run(INSTANCE_NAME)
        inst_run.font.size = Pt(16)
        inst_run.bold = True
        inst.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_paragraph()

        for metric, label in METRICS:
            p = doc.add_paragraph()
            run = p.add_run(label)
            run.bold = True
            run.font.size = Pt(12)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            img_data = fetch_chart(metric, INSTANCE_NAME)
            img_path = f"/tmp/{metric}.png"
            with open(img_path, "wb") as f:
                f.write(img_data)
            doc.add_picture(img_path, width=Inches(6.5))
            doc.add_paragraph()

        out_path = "/tmp/SCL_Report.docx"
        doc.save(out_path)
        s3.upload_file(out_path, bucket, key)

        return {"statusCode": 200, "body": f"Uploaded to s3://{bucket}/{key}"}

    except Exception as e:
        import traceback
        return {
            "statusCode": 500,
            "body": json.dumps({
                "errorType": str(type(e)),
                "errorMessage": str(e),
                "trace": traceback.format_exc()
            })
        }

