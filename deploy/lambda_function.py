import sys
import json
import boto3
import os
import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

# AWS clients
cw = boto3.client("cloudwatch", region_name="ap-south-1")
s3 = boto3.client("s3", region_name="ap-south-1")

# Constants for S3 location of your logos
LOGO_BUCKET     = "test-resource-util"
LOGO_PREFIX     = "lambda/Logos"
LEFT_LOGO_NAME  = "Hathi-Cement.png"
RIGHT_LOGO_NAME = "Sapphire.jpg"

INSTANCE_ID   = "i-0bf7a1cdc65f2a2fe"
INSTANCE_NAME = "PRD_S4H_SCL_APP"
REGION        = "ap-south-1"
METRICS = [
    ("NetworkIn",        "Network In"),
    ("NetworkOut",       "Network Out"),
    ("mem_used_percent", "Memory Utilization"),
    ("CPUUtilization",   "CPU Utilization")
]

def fetch_chart(metric, instance_label):
    widget = {
        "metrics": [["AWS/EC2", metric, "InstanceId", INSTANCE_ID]],
        "stat": "Average",
        "view": "timeSeries",
        "stacked": False,
        "region": REGION,
        "start": "-PT24H",
        "end": "P0D",
        "width": 1100,
        "height": 300,
        "title": instance_label,
        "setPeriodToTimeRange": True,
        "legend": {"position": "bottom"}
    }
    resp = cw.get_metric_widget_image(MetricWidget=json.dumps(widget))
    return resp["MetricWidgetImage"]

def download_logos_to_tmp():
    """
    Download both logos from S3 into /tmp so they can be embedded in the header.
    """
    for name in (LEFT_LOGO_NAME, RIGHT_LOGO_NAME):
        key = f"{LOGO_PREFIX}/{name}"
        dest = f"/tmp/{name}"
        s3.download_file(LOGO_BUCKET, key, dest)

def set_header_footer(doc):
    """
    Build a two-cell header table: left Hathi logo, right Sapphire logo,
    each cell exactly 3.25" wide within a 6.5" table, logos same size.
    """
    section = doc.sections[0]
    header = section.header

    # 1×2 table spanning 6.5" (full page width minus margins)
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Force each column to half width (3.25")
    cols = table.columns
    for col in cols:
        col.width = Inches(3.25)
        for cell in col.cells:
            cell.width = Inches(3.25)

    # Remove all borders
    tbl = table._tbl
    for tc in tbl.iter_tcs():
        tc.tcPr.append(OxmlElement("w:tcBorders"))

    # Both logos will be 1" wide for symmetry
    logo_width = Inches(1.0)

    # Left cell: Hathi logo, flush left
    left_cell = table.rows[0].cells[0]
    lp = left_cell.paragraphs[0]
    lp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    lp.add_run().add_picture(f"/tmp/{LEFT_LOGO_NAME}", width=logo_width)

    # Right cell: Sapphire logo, flush right
    right_cell = table.rows[0].cells[1]
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    rp.add_run().add_picture(f"/tmp/{RIGHT_LOGO_NAME}", width=logo_width)

    # Footer unchanged
    footer = section.footer.paragraphs[0]
    footer.text = f"Generated on {datetime.date.today():%d-%m-%Y}"
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer.runs[0].font.size = Pt(8)

def lambda_handler(event, context):
    try:
        # Download the logos into /tmp
        download_logos_to_tmp()

        today = datetime.date.today().strftime("%d-%m-%Y")
        bucket = event.get("ReportBucket", LOGO_BUCKET)
        key = f"Test/{datetime.date.today():%Y-%m-%d}/SCL_Report.docx"

        doc = Document()

        # Insert header with symmetric logos
        set_header_footer(doc)

        # Title and metadata in the body
        title = doc.add_paragraph()
        run = title.add_run("SCL Resource Utilization")
        run.font.size = Pt(20)
        run.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        date_para = doc.add_paragraph(f"Date: {today}")
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()

        inst = doc.add_paragraph()
        inst_run = inst.add_run(INSTANCE_NAME)
        inst_run.font.size = Pt(16)
        inst_run.bold = True
        inst.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_paragraph()

        # Embed each metric chart
        for metric, label in METRICS:
            p = doc.add_paragraph()
            run = p.add_run(label)
            run.bold = True
            run.font.size = Pt(12)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            img_data = fetch_chart(metric, INSTANCE_NAME)
            img_path = f"/tmp/{metric}.png"
            with open(img_path, "wb") as f:
                f.write(img_data)
            doc.add_picture(img_path, width=Inches(6.5))
            doc.add_paragraph()

        # Save and upload
        out_path = "/tmp/SCL_Report.docx"
        doc.save(out_path)
        s3.upload_file(out_path, bucket, key)

        return {
            "statusCode": 200,
            "body": f"Uploaded to s3://{bucket}/{key}"
        }

    except Exception as e:
        import traceback
        return {
            "statusCode": 500,
            "body": json.dumps({
                "errorType": str(type(e)),
                "errorMessage": str(e),
                "trace": traceback.format_exc()
            })
        }

