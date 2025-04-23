
import sys
import json
import boto3
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Add the layer path (required only in local test, Lambda handles this at runtime)
# sys.path.insert(0, "/opt/python")

cw = boto3.client("cloudwatch", region_name="ap-south-1")
s3 = boto3.client("s3", region_name="ap-south-1")

INSTANCE_ID = "i-0bf7a1cdc65f2a2fe"
INSTANCE_NAME = "PRD_S4H_SCL_APP"
REGION = "ap-south-1"

METRICS = [
    ("NetworkIn",        "Network In"),
    ("NetworkOut",       "Network Out"),
    ("mem_used_percent", "Memory Utilization"),
    ("CPUUtilization",   "CPU Utilization")
]

def fetch_chart(metric, title):
    widget = {
        "metrics": [["AWS/EC2", metric, "InstanceId", INSTANCE_ID]],
        "start": "-PT24H",
        "end": "P0D",
        "view": "timeSeries",
        "stacked": False,
        "stat": "Average",
        "width": 800,
        "height": 300,
        "title": title
    }
    return cw.get_metric_widget_image(MetricWidget=json.dumps(widget))["MetricWidgetImage"]

def lambda_handler(event, context):
    today = datetime.date.today().strftime("%d-%m-%Y")
    bucket = event.get("ReportBucket", "test-resource-util")
    key = f"Test/{datetime.date.today():%Y-%m-%d}/SCL_Report.docx"

    # Create DOCX
    doc = Document()

    # Title
    title = doc.add_heading("SCL Resource Utilization", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.runs[0]
    run.font.size = Pt(20)
    run.bold = True

    doc.add_paragraph(f"Date: {today}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    # Instance Section
    doc.add_heading(INSTANCE_NAME, level=1).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 2×2 Table
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"

    coords = [(0, 0), (0, 1), (1, 0), (1, 1)]
    for (metric, label), (r, c) in zip(METRICS, coords):
        cell = table.rows[r].cells[c]
        cell.paragraphs[0].add_run(label).bold = True
        img = fetch_chart(metric, label)
        path = f"/tmp/{metric}.png"
        with open(path, "wb") as f:
            f.write(img)
        cell.add_paragraph().add_run().add_picture(path, width=Inches(3.0))

    # Save to /tmp
    output_path = "/tmp/SCL_Report.docx"
    doc.save(output_path)

    # Upload to S3
    s3.upload_file(output_path, bucket, key)

    return {
        "statusCode": 200,
        "body": f"Uploaded to s3://{bucket}/{key}"
    }

